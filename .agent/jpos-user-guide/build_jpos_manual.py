from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Github\POS")
ASSETS = ROOT / ".agent" / "jpos-user-guide" / "screenshots"
OUTPUT = ROOT / "docs" / "Huong_dan_su_dung_JPOS.docx"

FONT = "Times New Roman"
BODY_SIZE = 12
HEADING_SIZE = 14
INK = "1F2937"
ORANGE = "D85B1F"
DARK_ORANGE = "A63F12"
MUTED = "5F6B7A"
LIGHT_ORANGE = "FFF3E8"
LIGHT_BLUE = "EFF6FF"
LIGHT_GREEN = "ECFDF5"
LIGHT_AMBER = "FFFBEB"
LIGHT_RED = "FEF2F2"
BORDER = "D9DEE7"


def set_run_font(run, size=BODY_SIZE, bold=None, italic=None, color=INK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_paragraph(paragraph, fill: str, border: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "6")
        edge.set(qn("w:space"), "6")
        edge.set(qn("w:color"), border)
        pBdr.append(edge)
    pPr.append(pBdr)


def add_field(paragraph, instruction: str, display="1"):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])
    set_run_font(run, size=10, color=MUTED)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (DARK_ORANGE, 16, 8),
        "Heading 2": (ORANGE, 13, 6),
        "Heading 3": (INK, 10, 4),
    }
    for name, (color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(HEADING_SIZE)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(BODY_SIZE)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(BODY_SIZE)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption._element.rPr.rFonts.set(qn("w:cs"), FONT)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.15

    # Patch numbering definitions used by Word built-in list styles.
    numbering = doc.part.numbering_part.element
    for lvl in numbering.xpath(".//w:lvl[@w:ilvl='0']"):
        pPr = lvl.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            lvl.append(pPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")


def configure_sections(doc: Document):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.78)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.38)
        section.footer_distance = Inches(0.38)

        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        hr = hp.add_run("JPOS - Sổ tay người dùng")
        set_run_font(hr, size=9, color=MUTED, italic=True)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_after = Pt(0)
        r = fp.add_run("Trang ")
        set_run_font(r, size=10, color=MUTED)
        add_field(fp, "PAGE")


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=HEADING_SIZE, bold=True, color=DARK_ORANGE)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=BODY_SIZE, italic=True, color=MUTED)
    return p


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def _new_step_numbering(doc: Document) -> int:
    """Create a real single-level numbered list that restarts at 1."""
    numbering = doc.part.numbering_part.element
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    num_id = max(num_ids, default=0) + 1

    base_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    base_num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def add_steps(doc, steps: Iterable[tuple[str, str]]):
    for index, (action, detail) in enumerate(steps, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.375)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        step = p.add_run(f"Bước {index}. ")
        set_run_font(step, bold=True, color=DARK_ORANGE)
        r1 = p.add_run(action)
        set_run_font(r1, bold=True, color=DARK_ORANGE)
        if detail:
            r2 = p.add_run(f" {detail}")
            set_run_font(r2)


def add_note(doc, label: str, text: str, tone="blue"):
    tones = {
        "blue": (LIGHT_BLUE, "93C5FD", "1D4ED8"),
        "green": (LIGHT_GREEN, "86EFAC", "166534"),
        "amber": (LIGHT_AMBER, "FCD34D", "92400E"),
        "red": (LIGHT_RED, "FCA5A5", "991B1B"),
        "orange": (LIGHT_ORANGE, "FDBA74", DARK_ORANGE),
    }
    fill, border, color = tones[tone]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, fill, border)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, color=color)
    return p


def add_figure(doc, filename: str, caption: str, width=6.45):
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    inline._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
        cell._tc.tcPr[-1].set(qn("w:fill"), LIGHT_ORANGE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, bold=True, color=DARK_ORANGE)
    trPr = hdr._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    trPr.append(repeat)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            set_run_font(r)
    set_table_geometry(table, widths_dxa)
    return table


def chapter_break(doc):
    # Heading 1 carries page_break_before, which avoids blank pages when the
    # previous chapter happens to end exactly at the page boundary.
    return None


def build():
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    core = doc.core_properties
    core.title = "Hướng dẫn sử dụng phần mềm JPOS"
    core.subject = "Sổ tay người dùng chi tiết"
    core.author = "Đội dự án JPOS"
    core.keywords = "JPOS, POS, hướng dẫn sử dụng, bán hàng, thành viên, kết ca"

    # Cover
    add_title(doc, "HƯỚNG DẪN SỬ DỤNG PHẦN MỀM JPOS")
    add_subtitle(doc, "Sổ tay thao tác đầy đủ cho nhân viên quầy, quản lý ca và người quản trị cấu hình")
    add_figure(doc, "01_ban_hang_tong_quan.png", "Giao diện bán hàng chính của JPOS (dữ liệu minh họa)", width=6.25)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    for line in (
        "Phiên bản phần mềm tham chiếu: 0.1.0",
        "Ngày biên soạn: 12/08/2026",
        "Phạm vi: ứng dụng quầy JPOS và màn hình khách hàng",
    ):
        r = p.add_run(line + "\n")
        set_run_font(r, color=MUTED)
    add_note(doc, "Lưu ý về hình ảnh", "Ảnh trong tài liệu được chụp trực tiếp từ giao diện JPOS bằng dữ liệu mẫu an toàn. Tên cửa hàng, khách hàng, số tiền và mã đơn trên máy thật sẽ khác.", "orange")

    chapter_break(doc)
    add_heading(doc, "MỤC LỤC", 1)
    toc_items = [
        "1. Giới thiệu JPOS và phạm vi chức năng",
        "2. Chuẩn bị trước khi sử dụng",
        "3. Kích hoạt thiết bị, đăng nhập và điều hướng",
        "4. Bán hàng và quản lý giỏ hàng",
        "5. Thanh toán, in biên lai và in vé",
        "6. Lịch sử và xử lý đơn hàng",
        "7. Kết ca và in báo cáo",
        "8. Quản lý thành viên",
        "9. Cài đặt JPOS",
        "10. Màn hình khách hàng",
        "11. Trạng thái, cảnh báo và xử lý sự cố",
        "12. Quy tắc vận hành an toàn và phụ lục tra cứu nhanh",
    ]
    add_bullets(doc, toc_items)
    add_note(doc, "Cách đọc", "Mỗi mục gồm: mục đích, thành phần trên trang, các bước thao tác, kết quả mong đợi và lưu ý. Tên nút được in đậm đúng như trên giao diện.", "blue")

    chapter_break(doc)
    add_heading(doc, "1. GIỚI THIỆU JPOS VÀ PHẠM VI CHỨC NĂNG", 1)
    add_body(doc, "JPOS là phần mềm điểm bán hàng dành cho quầy vận hành Joy World/Joyaland. Ứng dụng ưu tiên thao tác nhanh tại quầy, ghi nhận thanh toán tại chỗ và đồng bộ đơn hàng nền với hệ thống trung tâm.")
    add_heading(doc, "1.1. Nhóm chức năng chính", 2)
    add_bullets(doc, [
        "Kích hoạt máy POS bằng mã một lần do quản trị viên tạo trên JPULSE.",
        "Đăng nhập bằng email, tên đăng nhập hoặc số điện thoại; chọn cửa hàng khi tài khoản có nhiều điểm bán.",
        "Tìm kiếm, lọc, đồng bộ sản phẩm; thêm vé và hàng lưu niệm vào giỏ.",
        "Gắn thành viên vào đơn bằng số điện thoại hoặc đầu đọc thẻ.",
        "Thanh toán tiền mặt hoặc chuyển khoản/QR; hỗ trợ QR PayOS và QR tài khoản cố định theo cấu hình.",
        "Tự động in biên lai và vé sau thanh toán; in lại từ lịch sử đơn.",
        "Tra cứu đơn trong ngày, xem chi tiết, lọc trạng thái và thử đồng bộ lại đơn lỗi.",
        "Kết ca theo ngày hoặc khoảng ca; tổng hợp sản phẩm, tiền mặt và chuyển khoản.",
        "Tra cứu, đăng ký, nạp thẻ, xem biến động/vé/thẻ và nạp bù thành viên theo quyền.",
        "Cấu hình biên lai, vé, thanh toán và playlist quảng cáo màn hình khách.",
        "Màn hình khách hiển thị quảng cáo, giỏ hàng, QR, trạng thái thanh toán và thông tin thành viên an toàn.",
    ])
    add_heading(doc, "1.2. Luồng dữ liệu đơn hàng", 2)
    add_table(doc, ["Trạng thái", "Ý nghĩa vận hành"], [
        ["DRAFT", "Đơn nháp đang được chuẩn bị hoặc giỏ hàng đang hình thành."],
        ["LOCAL_PAID", "Đã ghi nhận thanh toán cục bộ; chờ tiến trình đồng bộ nền."],
        ["SYNCING", "Hệ thống đang gửi đơn/thanh toán lên dịch vụ trung tâm."],
        ["SYNC_SUCCESS", "Đồng bộ thành công; có thể đã nhận mã đơn HK."],
        ["SYNC_FAILED", "Đồng bộ thất bại; cần kiểm tra lỗi và dùng Thử lại."],
    ], [1900, 7460])
    add_note(doc, "Quan trọng", "Thanh toán thành công tại quầy và đồng bộ thành công là hai mốc khác nhau. Không thu tiền khách lần thứ hai chỉ vì đơn đang chờ hoặc lỗi đồng bộ.", "red")

    chapter_break(doc)
    add_heading(doc, "2. CHUẨN BỊ TRƯỚC KHI SỬ DỤNG", 1)
    add_heading(doc, "2.1. Điều kiện bắt buộc", 2)
    add_bullets(doc, [
        "Máy POS đã cài JPOS và có kết nối mạng cho lần kích hoạt/xác minh đầu tiên.",
        "Có mã kích hoạt 8 số còn hiệu lực do quản trị viên tạo trên JPULSE.",
        "Có tài khoản người dùng đang hoạt động và đã được gán ít nhất một kho/cửa hàng loại STORE.",
        "Máy in nhiệt và driver Windows đã cài nếu cần in biên lai/vé; đặt đúng máy in mặc định.",
        "Đầu đọc thẻ được kết nối nếu quầy dùng tra cứu thành viên bằng thẻ.",
        "Màn hình phụ được kết nối nếu sử dụng màn hình khách hàng.",
    ])
    add_heading(doc, "2.2. Kiểm tra đầu ca", 2)
    add_steps(doc, [
        ("Kiểm tra mạng.", "Đảm bảo máy truy cập được dịch vụ nội bộ và JPULSE."),
        ("Kiểm tra máy in.", "In thử trong thẻ Cài đặt > In biên lai và Cài đặt > In vé."),
        ("Kiểm tra màn hình khách.", "Màn hình chờ phải hiển thị quảng cáo mặc định hoặc playlist của cửa hàng."),
        ("Đồng bộ sản phẩm.", "Tại trang Bán hàng, bấm biểu tượng mũi tên vòng tròn nếu danh mục chưa mới."),
        ("Kiểm tra cửa hàng.", "Tên cửa hàng hiển thị phía trên lời chào phải đúng với quầy đang vận hành."),
    ])
    add_note(doc, "Quyền offline", "Sau một lần xác minh hợp lệ, thiết bị có thể dùng quyền và cấu hình đã cache trong thời gian giới hạn. Nếu quá 8 giờ chưa kết nối lại JPULSE, JPOS sẽ chặn tiếp tục để bảo đảm quyền truy cập còn hợp lệ.", "amber")

    chapter_break(doc)
    add_heading(doc, "3. KÍCH HOẠT THIẾT BỊ, ĐĂNG NHẬP VÀ ĐIỀU HƯỚNG", 1)
    add_heading(doc, "3.1. Kích hoạt máy POS", 2)
    add_figure(doc, "00_kich_hoat_thiet_bi.png", "Màn hình Kích hoạt máy POS - nhập tên quầy và mã kích hoạt 8 số")
    add_body(doc, "Trang này xuất hiện khi thiết bị chưa có thông tin tin cậy hoặc phiên kích hoạt trước đã bị xóa/chặn.")
    add_steps(doc, [
        ("Nhập Tên quầy.", "Dùng tên dễ nhận biết, ví dụ Quầy vé tầng 1 hoặc POS-02."),
        ("Nhập Mã kích hoạt.", "Nhập đủ 8 chữ số do quản trị viên tạo trên JPULSE."),
        ("Bấm Kích hoạt thiết bị.", "Nút chỉ khả dụng khi tên quầy có nội dung và mã đủ 8 số."),
        ("Chờ xác minh.", "Khi thành công, thiết bị nhận cửa hàng, quyền và cấu hình từ JPULSE rồi chuyển đến đăng nhập."),
    ])
    add_note(doc, "Thời hạn mã", "Mã kích hoạt chỉ dùng một lần và hết hạn sau 10 phút. Nếu hết hạn, yêu cầu quản trị viên tạo mã mới; không thử lặp lại mã cũ.", "amber")
    add_heading(doc, "3.2. Đăng nhập", 2)
    add_figure(doc, "00_dang_nhap.png", "Màn hình đăng nhập JPOS")
    add_steps(doc, [
        ("Nhập định danh.", "Có thể dùng email, tên đăng nhập hoặc số điện thoại nhân viên."),
        ("Nhập mật khẩu.", "Bấm biểu tượng mắt để hiện/ẩn nội dung mật khẩu khi cần kiểm tra."),
        ("Bấm Đăng nhập.", "JPOS xác thực tài khoản, vai trò, quyền và danh sách cửa hàng được phép truy cập."),
        ("Chọn cửa hàng nếu được yêu cầu.", "Bấm đúng thẻ cửa hàng cho phiên POS hiện tại. Nếu chỉ có một cửa hàng hợp lệ, hệ thống tự chọn."),
    ])
    add_note(doc, "Không đăng nhập được", "Kiểm tra định danh/mật khẩu, mạng, trạng thái tài khoản, quyền POS và việc tài khoản đã được gán cửa hàng loại STORE. Sau nhiều lần sai, hãy dừng thử và liên hệ quản trị viên.", "red")
    add_heading(doc, "3.3. Thanh điều hướng bên trái", 2)
    add_table(doc, ["Biểu tượng/mục", "Chức năng"], [
        ["Bán hàng", "Danh mục sản phẩm, giỏ hàng và thanh toán."],
        ["Đơn hàng", "Đơn trong ngày, bộ lọc, chi tiết và đồng bộ lại."],
        ["Kết ca", "Tổng hợp sản phẩm bán ra và doanh thu theo kỳ."],
        ["Thành viên", "Tra cứu, đăng ký, nạp thẻ, biến động và nạp bù."],
        ["Cài đặt", "Biên lai, vé, thanh toán và quảng cáo."],
        ["Đăng xuất", "Mở hộp xác nhận kết thúc phiên người dùng."],
    ], [2400, 6960])
    add_heading(doc, "3.4. Đăng xuất", 2)
    add_figure(doc, "16_dang_xuat.png", "Hộp thoại xác nhận đăng xuất")
    add_steps(doc, [
        ("Bấm biểu tượng Đăng xuất.", "Biểu tượng nằm cuối thanh bên trái."),
        ("Chọn Hủy nếu bấm nhầm.", "Có thể nhấn Esc hoặc bấm ngoài hộp thoại để đóng."),
        ("Chọn Đăng xuất để xác nhận.", "Phiên đăng nhập và lựa chọn cửa hàng của người dùng sẽ được xóa."),
    ])
    add_note(doc, "Trước khi đăng xuất", "Hoàn tất hoặc hủy phiên QR đang hoạt động, xử lý các cảnh báo khôi phục thanh toán và bảo đảm không còn khách đang chờ nhận biên lai/vé.", "amber")

    chapter_break(doc)
    add_heading(doc, "4. BÁN HÀNG VÀ QUẢN LÝ GIỎ HÀNG", 1)
    add_heading(doc, "4.1. Thành phần trang Bán hàng", 2)
    add_figure(doc, "01_ban_hang_tong_quan.png", "Trang Bán hàng gồm thanh điều hướng, tìm kiếm/danh mục, lưới sản phẩm và giỏ hàng")
    add_bullets(doc, [
        "Góc trên trái: tên cửa hàng và lời chào theo tên nhân viên.",
        "Ô Tìm tên, nhóm hoặc mã vạch: tìm theo tên sản phẩm, nhóm hoặc chuỗi mã liên quan.",
        "Nút Đồng bộ sản phẩm: tải catalog mới nhất rồi nạp lại danh mục tại quầy.",
        "Nút loại sản phẩm: Vé lượt và Sản phẩm lưu niệm.",
        "Dải nhóm con: Tất cả, Quy đổi vé Online, Combo, Vé OTA, Quy đổi eVoucher... tùy dữ liệu.",
        "Thẻ sản phẩm: nhóm, tên, giá sau thuế/miễn phí và chỉ báo cấu hình số vé.",
        "Khung Đơn hàng hiện tại: thành viên, dòng hàng, số lượng, tổng tiền và nút Thanh toán.",
    ])
    add_heading(doc, "4.2. Tìm, lọc và đồng bộ sản phẩm", 2)
    add_steps(doc, [
        ("Chọn loại sản phẩm.", "Bấm biểu tượng vé hoặc quà để chuyển danh mục chính."),
        ("Chọn nhóm con.", "Bấm Tất cả hoặc nhóm nghiệp vụ phù hợp; số bên cạnh cho biết số mặt hàng."),
        ("Nhập nội dung tìm kiếm.", "Nhập toàn bộ hoặc một phần tên, nhóm hoặc mã vạch."),
        ("Bấm Đồng bộ sản phẩm khi cần.", "Chờ biểu tượng ngừng quay rồi kiểm tra lại số lượng/giá."),
    ])
    add_note(doc, "Giá bán", "JPOS ưu tiên giá sau thuế của sản phẩm. Nếu sản phẩm hiển thị 0 đ, xác nhận đây thực sự là vé/quà miễn phí trước khi thêm vào đơn.", "amber")
    add_heading(doc, "4.3. Thêm sản phẩm vào giỏ", 2)
    add_figure(doc, "02_gio_hang.png", "Giỏ hàng sau khi bấm các thẻ sản phẩm")
    add_steps(doc, [
        ("Bấm vào thẻ sản phẩm.", "Sản phẩm được thêm với số lượng 1; đơn nháp có thể được chuẩn bị tự động."),
        ("Bấm dấu + để tăng số lượng.", "Tổng dòng hàng và tổng cộng cập nhật ngay."),
        ("Bấm dấu - để giảm số lượng.", "Khi giảm về 0, dòng hàng được loại khỏi giỏ."),
        ("Bấm biểu tượng thùng rác trên dòng.", "Xóa riêng mặt hàng đang chọn."),
        ("Bấm thùng rác ở tiêu đề giỏ.", "Xóa toàn bộ giỏ; không khả dụng khi phiên thanh toán đang khóa."),
    ])
    add_heading(doc, "4.4. Gắn thành viên vào đơn", 2)
    add_body(doc, "Chỉ nên gắn thành viên sau khi giỏ có ít nhất một sản phẩm. JPOS hỗ trợ hai cách:")
    add_steps(doc, [
        ("Theo số điện thoại.", "Nhập số vào ô Số điện thoại thành viên, sau đó bấm Tìm."),
        ("Theo thẻ.", "Để trống số điện thoại, bấm Đọc thẻ rồi đưa thẻ vào đầu đọc. Bấm Hủy nếu cần dừng chờ."),
        ("Kiểm tra thẻ thành viên đã gắn.", "Tên và mã/số điện thoại/hạng thành viên xuất hiện trong khung xanh phía trên giỏ."),
        ("Gỡ thành viên nếu chọn nhầm.", "Bấm dấu X trên thẻ thành viên trước khi thanh toán."),
    ])
    add_note(doc, "Không đọc được thẻ", "Kiểm tra cáp/driver đầu đọc, đặt thẻ đúng vị trí và thử lại. Có thể chuyển sang tìm bằng số điện thoại để tiếp tục phục vụ khách.", "blue")

    chapter_break(doc)
    add_heading(doc, "5. THANH TOÁN, IN BIÊN LAI VÀ IN VÉ", 1)
    add_heading(doc, "5.1. Mở cửa sổ thanh toán", 2)
    add_steps(doc, [
        ("Kiểm tra giỏ hàng.", "Đối chiếu tên hàng, số lượng, thành viên và tổng cộng."),
        ("Bấm Thanh toán.", "Nút chỉ bật khi giỏ có sản phẩm và hệ thống không bận xử lý giao dịch khác."),
        ("Chọn ngôn ngữ biên lai.", "Chọn Tiếng Việt, Tiếng Anh hoặc Tiếng Trung trước khi xác nhận."),
        ("Chọn phương thức.", "Dùng Tiền mặt hoặc Chuyển khoản theo thực tế khách trả."),
    ])
    add_heading(doc, "5.2. Thanh toán tiền mặt", 2)
    add_figure(doc, "03_thanh_toan_tien_mat.png", "Cửa sổ thanh toán tiền mặt với các mệnh giá VND")
    add_steps(doc, [
        ("Chọn Tiền mặt.", "Thẻ Tiền mặt có viền cam và dấu chọn."),
        ("Bấm từng tờ tiền khách đưa.", "Mỗi lần bấm cộng thêm 1 tờ tương ứng: 500.000, 200.000, 100.000, 50.000, 20.000, 10.000, 5.000, 2.000 hoặc 1.000 đ."),
        ("Sửa số tờ nhập nhầm.", "Bấm huy hiệu số tờ/dấu trừ trên mệnh giá để giảm từng tờ."),
        ("Theo dõi Còn thiếu và Cần thối.", "Nút xác nhận chỉ bật khi tiền đã nhận đủ."),
        ("Xác nhận thanh toán.", "Đọc lại tiền khách đưa và tiền thối, sau đó bấm nút xác nhận cuối hộp thoại."),
    ])
    add_note(doc, "Kiểm soát tiền mặt", "Không bấm xác nhận trước khi đã nhận và đếm tiền. Luôn trả đúng số Cần thối hiển thị trên màn hình.", "red")
    add_heading(doc, "5.3. Thanh toán chuyển khoản/QR", 2)
    add_figure(doc, "04_thanh_toan_chuyen_khoan.png", "Chọn Chuyển khoản và nút Tạo mã thanh toán")
    add_steps(doc, [
        ("Chọn Chuyển khoản.", "JPOS chuyển từ bảng tiền mặt sang quy trình QR."),
        ("Bấm Tạo mã thanh toán.", "Hệ thống tạo phiên PayOS hoặc QR cố định theo cấu hình cửa hàng."),
        ("Yêu cầu khách quét QR trên màn hình khách.", "Đối chiếu đúng số tiền và nội dung chuyển khoản."),
        ("Chờ trạng thái thành công.", "Với PayOS, hệ thống theo dõi và tự hoàn tất khi nhận xác nhận hợp lệ."),
        ("Với QR cố định, kiểm tra giao dịch thực tế.", "Nhân viên phải kiểm tra tài khoản ngân hàng và bấm xác nhận thủ công."),
        ("Nếu QR hết hạn/hủy/lỗi.", "Không sửa giỏ khi phiên đang khóa; hủy hoặc tạo lại theo nút trạng thái trên màn hình."),
    ])
    add_note(doc, "QR cố định", "Đơn xác nhận thủ công bằng QR tài khoản cố định được đánh dấu Chưa được xác nhận thanh toán trong lịch sử. Đây là chỉ báo để quản lý đối soát lại giao dịch ngân hàng.", "amber")
    add_heading(doc, "5.4. Voucher", 2)
    add_body(doc, "Nút Có mã voucher? mở ô nhập mã và nút quét camera. Tuy nhiên, phiên bản hiện tại chưa có OpenAPI xác thực/khấu trừ voucher và chức năng quét camera vẫn là mục dự kiến. Khi bấm Áp dụng, JPOS sẽ cảnh báo chưa thể áp dụng.")
    add_note(doc, "Không hứa giảm giá", "Không xác nhận với khách rằng voucher đã được trừ nếu tổng thanh toán chưa thay đổi và hệ thống chưa báo áp dụng thành công.", "red")
    add_heading(doc, "5.5. Sau khi thanh toán thành công", 2)
    add_steps(doc, [
        ("Chờ thông báo thành công.", "Đơn được lưu và tiến trình đồng bộ nền bắt đầu."),
        ("Nhận biên lai/vé từ máy in.", "Biên lai được in tự động; vé được in sau biên lai nếu đã bật Tự động in vé."),
        ("Kiểm tra đủ số vé.", "Mỗi mã vé là một trang/mảnh riêng; số vé phụ thuộc snapshot của sản phẩm trong đơn."),
        ("Nếu in lỗi, không thu tiền lại.", "Mở Lịch sử đơn, chọn đúng đơn và bấm In biên lai; kiểm tra cấu hình máy in/vé."),
        ("Theo dõi mã và trạng thái đơn.", "LOCAL_PAID/SYNCING sẽ được làm mới định kỳ; SYNC_SUCCESS là hoàn tất đồng bộ."),
    ])
    add_heading(doc, "5.6. Khôi phục giao dịch dang dở", 2)
    add_body(doc, "Nếu ứng dụng đóng hoặc giao diện lỗi trong lúc thanh toán, JPOS có thể hiển thị thông báo khôi phục từ nhật ký checkout. Đọc kỹ mã đơn, trạng thái và hành động được gợi ý; mở lịch sử đơn để xác minh trước khi thử thanh toán lại.")
    add_note(doc, "Nguyên tắc chống thanh toán trùng", "Khi không chắc một giao dịch đã thành công hay chưa, kiểm tra lịch sử đơn và tài khoản ngân hàng/PayOS trước. Không tạo giao dịch thứ hai chỉ để thử.", "red")

    chapter_break(doc)
    add_heading(doc, "6. LỊCH SỬ VÀ XỬ LÝ ĐƠN HÀNG", 1)
    add_heading(doc, "6.1. Thành phần trang Lịch sử đơn hàng", 2)
    add_figure(doc, "05_lich_su_don_hang.png", "Trang Lịch sử đơn hàng - KPI, bộ lọc và các thẻ đơn trong ngày")
    add_bullets(doc, [
        "KPI Tổng đơn hàng: số đơn sau bộ lọc.",
        "KPI Doanh thu thành công: chỉ cộng đơn SYNC_SUCCESS.",
        "KPI Đã đồng bộ: số đơn SYNC_SUCCESS.",
        "KPI Chờ/Lỗi đồng bộ: LOCAL_PAID, SYNCING hoặc SYNC_FAILED.",
        "Bộ lọc Trạng thái: Tất cả, Đã đồng bộ, Chờ đồng bộ, Đang đồng bộ, Lỗi đồng bộ.",
        "Ô tìm kiếm: mã local/mã HK, tên sản phẩm, tên/số điện thoại/mã thành viên.",
        "Sắp xếp: Mới nhất, Cũ nhất, Giá giảm dần, Giá tăng dần.",
        "Danh sách chỉ hiển thị đơn của ngày hiện tại.",
    ])
    add_heading(doc, "6.2. Tra cứu đơn", 2)
    add_steps(doc, [
        ("Bấm Làm mới.", "Tải lại danh sách từ dịch vụ POS."),
        ("Chọn trạng thái.", "Thu hẹp nhóm đơn cần xử lý."),
        ("Nhập từ khóa.", "Có thể dùng một phần mã đơn, tên hàng hoặc thông tin khách."),
        ("Chọn cách sắp xếp.", "Ưu tiên Mới nhất khi phục vụ khách vừa thanh toán."),
        ("Bấm thẻ đơn.", "Mở cửa sổ Chi tiết đơn hàng."),
    ])
    add_heading(doc, "6.3. Chi tiết đơn và in lại", 2)
    add_figure(doc, "06_chi_tiet_don_hang.png", "Chi tiết đơn hàng - trạng thái, khách, sản phẩm, voucher, đồng bộ và In biên lai")
    add_bullets(doc, [
        "Mã local, trạng thái, thời gian, phương thức thanh toán và mã HK.",
        "Thông tin khách hàng hoặc thành viên nếu đơn có gắn thành viên.",
        "Danh sách sản phẩm, đơn giá, số lượng và thành tiền.",
        "Voucher/giảm giá nếu có; tạm tính và tổng cộng.",
        "Số lần thử đồng bộ, thời điểm đồng bộ và lỗi gần nhất.",
        "Nút In biên lai để gửi lại chứng từ tới máy in mặc định.",
    ])
    add_heading(doc, "6.4. Đồng bộ lại đơn lỗi", 2)
    add_steps(doc, [
        ("Lọc Lỗi đồng bộ.", "Tìm các đơn SYNC_FAILED."),
        ("Kiểm tra thanh toán đã nhận.", "Đối chiếu tiền mặt hoặc giao dịch chuyển khoản trước khi thao tác."),
        ("Bấm Thử lại.", "Có thể bấm trên thẻ đơn hoặc trong chi tiết đơn."),
        ("Chờ danh sách làm mới.", "Đơn được xếp lại cho tiến trình nền; trạng thái có thể chuyển sang SYNCING rồi SYNC_SUCCESS."),
        ("Nếu tiếp tục lỗi.", "Ghi lại mã local, lỗi gần nhất và số lần thử để gửi bộ phận hỗ trợ."),
    ])

    chapter_break(doc)
    add_heading(doc, "7. KẾT CA VÀ IN BÁO CÁO", 1)
    add_heading(doc, "7.1. Thành phần trang Kết ca", 2)
    add_figure(doc, "07_ket_ca.png", "Trang Kết ca với bộ lọc, KPI, bảng sản phẩm và doanh thu theo phương thức")
    add_bullets(doc, [
        "Thời gian báo cáo: Theo ngày hoặc Theo ca.",
        "Theo ngày: chọn một ngày lịch; phạm vi từ 00:00 đến trước 00:00 ngày kế tiếp.",
        "Theo ca: nhập giờ bắt đầu và kết thúc; thời lượng tối đa 48 giờ.",
        "Phạm vi nhân viên: Của tôi hoặc Tất cả (tùy quyền/dữ liệu dịch vụ).",
        "Nút Tải báo cáo: áp dụng bộ lọc và lấy dữ liệu mới.",
        "Nút Làm mới: tải lại báo cáo theo bộ lọc hiện tại.",
        "Nút In báo cáo: in phiếu kết ca theo dữ liệu đang hiển thị.",
    ])
    add_heading(doc, "7.2. Lập báo cáo theo ngày", 2)
    add_steps(doc, [
        ("Chọn Theo ngày.", "Nút được tô nổi khi đang hoạt động."),
        ("Chọn ngày.", "Dùng bộ chọn lịch; kiểm tra đúng múi giờ/ngày vận hành."),
        ("Chọn Của tôi hoặc Tất cả.", "Của tôi chỉ phản ánh tài khoản hiện tại."),
        ("Bấm Tải báo cáo.", "Chờ các KPI và bảng dữ liệu cập nhật."),
        ("Đối chiếu số liệu.", "Kiểm tra số đơn, số sản phẩm, tiền mặt, chuyển khoản và tổng doanh thu."),
        ("Bấm In báo cáo.", "Lưu kèm chứng từ giao ca theo quy trình nội bộ."),
    ])
    add_heading(doc, "7.3. Lập báo cáo theo ca", 2)
    add_steps(doc, [
        ("Chọn Theo ca.", "Hai trường thời gian bắt đầu/kết thúc xuất hiện."),
        ("Nhập thời gian bắt đầu.", "Dùng thời điểm nhân viên nhận ca."),
        ("Nhập thời gian kết thúc.", "Phải sau thời gian bắt đầu và không vượt 48 giờ."),
        ("Chọn phạm vi nhân viên và tải báo cáo.", "Nếu đổi bộ lọc sau khi tải, tải lại trước khi in."),
    ])
    add_note(doc, "Phạm vi số liệu", "Báo cáo chỉ tính các đơn đã thanh toán do dịch vụ trả về trong kỳ. Dòng cập nhật cuối trang cho biết thời điểm dữ liệu được tải.", "blue")

    chapter_break(doc)
    add_heading(doc, "8. QUẢN LÝ THÀNH VIÊN", 1)
    add_heading(doc, "8.1. Ba chế độ nghiệp vụ", 2)
    add_bullets(doc, [
        "Tra cứu: tìm thành viên và xem hồ sơ, nạp thẻ, biến động, vé/gói, thẻ.",
        "Đăng ký mới: tạo hồ sơ thành viên và có thể chọn sản phẩm trước khi chuyển sang thanh toán.",
        "Nạp bù: cộng hoặc trừ điểm trực tiếp; chỉ xuất hiện khi người dùng có quyền pos.members.compensate.",
    ])
    add_heading(doc, "8.2. Tra cứu thành viên", 2)
    add_figure(doc, "08_thanh_vien_tra_cuu.png", "Tra cứu theo số điện thoại và thẻ Hồ sơ thành viên")
    add_steps(doc, [
        ("Chọn Tra cứu.", "Chế độ mặc định khi mở trang Thành viên."),
        ("Chọn Số điện thoại hoặc Mã thẻ.", "Nếu chọn Mã thẻ, có thể đọc thẻ hoặc nhập mã thủ công."),
        ("Nhập thông tin và bấm Tra cứu thành viên.", "Hệ thống lấy dữ liệu mới nhất từ OpenAPI."),
        ("Kiểm tra đúng khách.", "Đối chiếu họ tên, điện thoại, mã thẻ, hạng và UID/MID nếu cần."),
        ("Bấm Xóa phiên tra cứu.", "Xóa thông tin hiện tại trước khi phục vụ khách tiếp theo."),
    ])
    add_heading(doc, "8.3. Các thẻ chi tiết thành viên", 2)
    add_table(doc, ["Thẻ", "Nội dung và thao tác"], [
        ["Hồ sơ", "Số dư tài khoản VND, điểm thưởng, tổng khả dụng, điểm tích lũy, điện thoại, mã thẻ, ngày sinh, giới tính và email."],
        ["Nạp thẻ", "Tải catalog gói điểm, tìm gói, chọn gói, xem điểm gốc/thưởng, tiếp tục thanh toán tiền mặt hoặc QR và thử lại đúng đơn khi OpenAPI lỗi."],
        ["Biến động", "Lọc loại số dư, từ ngày/đến ngày, tải lại và chuyển trang lịch sử sử dụng/nạp."],
        ["Vé / gói", "Xem danh sách pass ticket/gói của thành viên và tải lại dữ liệu."],
        ["Thẻ", "Xem các thẻ gắn với thành viên, số IC/mã thành viên và ghi chú."],
    ], [1800, 7560])
    add_heading(doc, "8.4. Đăng ký thành viên mới", 2)
    add_figure(doc, "09_dang_ky_thanh_vien.png", "Biểu mẫu Đăng ký mới và danh mục sản phẩm có thể giữ trong đơn")
    add_steps(doc, [
        ("Chọn Đăng ký mới.", "Không thể chuyển nếu giỏ đang bị khóa bởi phiên thanh toán."),
        ("Chọn danh xưng.", "Ông, Bà, Anh, Chị hoặc dấu -; danh xưng được ghép vào họ tên."),
        ("Nhập Họ và tên và Số điện thoại.", "Hai trường có dấu * là bắt buộc."),
        ("Chọn giới tính.", "Nam hoặc Nữ theo dữ liệu khách cung cấp."),
        ("Nhập ngày sinh và email nếu có.", "Ngày/tháng/năm tách riêng; email phải đúng định dạng."),
        ("Bấm Đăng ký.", "Chờ OpenAPI xác nhận; khi thành công hồ sơ được lưu tại POS."),
        ("Tùy chọn chọn sản phẩm.", "Dùng nhóm Gói điểm, Vé hoặc Quà lưu niệm; bấm Thêm để giữ vào giỏ."),
        ("Bấm Đăng ký và thanh toán.", "Sau khi thành viên đã sẵn sàng, JPOS gắn thành viên vào đơn và quay về cửa sổ checkout."),
    ])
    add_heading(doc, "8.5. Nạp thẻ/gói điểm", 2)
    add_steps(doc, [
        ("Tra cứu đúng thành viên.", "Không nạp khi chưa đối chiếu khách và số dư."),
        ("Mở thẻ Nạp thẻ.", "Bấm Tải lại nếu catalog chưa có dữ liệu."),
        ("Tìm và chọn gói.", "Kiểm tra giá thanh toán, tổng điểm, điểm gốc và điểm thưởng."),
        ("Bấm Tiếp tục thanh toán gói.", "Chọn tiền mặt hoặc QR trong hộp thoại."),
        ("Hoàn tất thanh toán.", "Chờ OpenAPI xác nhận và ghi lại mã đơn từ xa nếu hiển thị."),
        ("Nếu lỗi từ xa.", "Dùng Thử lại đúng đơn hiện tại; không tạo một giao dịch gói mới khi chưa xác minh."),
    ])
    add_heading(doc, "8.6. Nạp bù hoặc trừ điểm", 2)
    add_figure(doc, "10_nap_bu_thanh_vien.png", "Màn hình Nạp bù - số điểm, lý do, cảnh báo kiểm toán và lịch sử")
    add_steps(doc, [
        ("Chọn Nạp bù và tra cứu khách.", "Đảm bảo đúng hồ sơ thành viên."),
        ("Nhập Số điểm điều chỉnh.", "Số dương để cộng; số âm để trừ; không nhập 0."),
        ("Chọn nhanh hoặc nhập lý do.", "Lý do tối thiểu 5 ký tự, tối đa 500 ký tự; nêu chứng từ/sự cố đã xác minh."),
        ("Bấm Kiểm tra cộng điểm hoặc Kiểm tra trừ điểm.", "Hộp xác nhận cho phép rà soát lần cuối."),
        ("Xác nhận thao tác.", "Thay đổi được ghi vĩnh viễn vào nhật ký kiểm toán."),
        ("Tải lại hồ sơ/lịch sử.", "Kiểm tra số dư mới và bản ghi biến động."),
    ])
    add_note(doc, "Nghiệp vụ nhạy cảm", "Không dùng Nạp bù để thay thế giao dịch bán gói thông thường. Chỉ thao tác khi có quyền, có lý do rõ ràng và theo quy trình phê duyệt của đơn vị.", "red")

    chapter_break(doc)
    add_heading(doc, "9. CÀI ĐẶT JPOS", 1)
    add_body(doc, "Bấm biểu tượng bánh răng ở cuối thanh bên. Thanh thẻ Cài đặt gồm In biên lai, In vé, Thanh toán và Quảng cáo. Khả năng chỉnh sửa phụ thuộc quyền người dùng và cửa hàng gắn với thiết bị.")
    add_heading(doc, "9.1. Cài đặt thanh toán", 2)
    add_figure(doc, "11_cai_dat_thanh_toan.png", "Cài đặt QR tài khoản cố định")
    add_bullets(doc, [
        "Bật QR cố định: cho phép dùng tài khoản ngân hàng được cấu hình làm phương án dự phòng.",
        "Chỉ sử dụng QR tài khoản cố định: không tạo PayOS và không tự xác nhận; nhân viên kiểm tra rồi xác nhận thủ công.",
        "Mã BIN ngân hàng: đúng 6 chữ số.",
        "Số tài khoản: từ 6 đến 19 chữ số.",
        "Tên chủ tài khoản: tối đa 50 ký tự, đúng tên hiển thị trên ứng dụng ngân hàng.",
        "Lưu cấu hình: chỉ khả dụng khi có quyền pos.settings.manage và dữ liệu tải xong.",
    ])
    add_steps(doc, [
        ("Chọn chế độ QR.", "Bật/tắt QR cố định và tùy chọn chỉ dùng QR cố định."),
        ("Nhập BIN, số tài khoản, tên chủ tài khoản.", "Kiểm tra chính xác từng chữ số."),
        ("Bấm Lưu cấu hình.", "Chờ thông báo lưu thành công trước khi rời trang."),
        ("Tạo giao dịch thử nội bộ theo quy trình.", "Quét QR và đối chiếu số tiền/nội dung; không dùng tài khoản khách thật để thử nếu chưa được phép."),
    ])
    add_heading(doc, "9.2. Cài đặt in biên lai", 2)
    add_figure(doc, "12_cai_dat_bien_lai.png", "Cài đặt biên lai với khổ giấy, thông tin cửa hàng và xem trước thực tế")
    add_bullets(doc, [
        "Khổ giấy: POS58, POS80 hoặc POS82; chọn đúng driver/hộp thoại in hệ thống.",
        "Nhận diện cửa hàng: tải logo, tên cửa hàng, địa chỉ, hotline; chỉnh rộng/cao/tương phản logo.",
        "Chủ đề: Mặc định, Quốc khánh 2/9 hoặc Tết; bật/tắt câu chủ đề, chỉnh nội dung và cỡ chữ.",
        "QR yêu cầu xuất hóa đơn: bật/tắt, chỉnh kích thước QR, cỡ tiêu đề và chú thích.",
        "Độ đậm nét in: chỉnh độc lập tên cửa hàng, địa chỉ, tiêu đề, thông tin đơn, hàng hóa, thuế, tổng tiền, QR, câu chủ đề và hậu mãi.",
        "Nội dung nâng cao: hiển thị logo/nhân viên/thuế/liên hệ, thuế mặc định cho đơn cũ, hậu mãi và lời cảm ơn.",
        "Xem trước thực tế: cập nhật theo cấu hình; In thử kiểm tra máy in mà không tạo đơn bán hàng.",
        "Khôi phục mặc định: đưa biểu mẫu về giá trị chuẩn; vẫn cần Lưu cấu hình để áp dụng.",
    ])
    add_steps(doc, [
        ("Chọn khổ giấy.", "Ưu tiên đúng vùng in thực của máy."),
        ("Nhập nhận diện cửa hàng.", "Kiểm tra dấu tiếng Việt trên bản xem trước."),
        ("Điều chỉnh chủ đề, QR và độ đậm.", "Tránh logo/QR quá nhỏ hoặc chữ quá mảnh trên máy in nhiệt."),
        ("Bấm In thử.", "Kiểm tra cắt giấy, độ đậm, QR và chiều rộng."),
        ("Bấm Lưu cấu hình.", "Chờ trạng thái Đã đồng bộ; cấu hình được đồng bộ hai chiều với JPULSE."),
    ])
    add_heading(doc, "9.3. Cài đặt in vé", 2)
    add_figure(doc, "13_cai_dat_ve.png", "Cài đặt vé nhiệt và xem trước vé có mã QR")
    add_bullets(doc, [
        "Khổ vé: POS58/POS80/POS82; chiều dài mỗi vé và kích thước QR.",
        "Nhận diện vé: logo, tên cửa hàng, tiêu đề vé, dòng giới thiệu, độ rộng/tương phản logo.",
        "Nội dung và cỡ chữ: tiêu đề, tên sản phẩm, nội dung phụ, độ đậm tên vé, hướng dẫn và lời nhắn cuối vé.",
        "Tự động in vé: in ngay sau biên lai khi thanh toán thành công.",
        "Tùy chọn hiển thị: logo, mã đơn, thời gian, giá và thứ tự Vé n/m.",
        "Mỗi mã vé được in thành một trang riêng; driver máy in cần cắt theo trang nếu máy không tự cắt.",
        "In thử và xem trước dùng dữ liệu mẫu; vé thật tạo mã QR duy nhất cho từng vé.",
    ])
    add_heading(doc, "9.4. Cài đặt quảng cáo màn hình khách", 2)
    add_figure(doc, "14_cai_dat_quang_cao.png", "Cài đặt playlist quảng cáo màn hình khách")
    add_bullets(doc, [
        "Playlist riêng theo cửa hàng và đồng bộ realtime với JPULSE.",
        "Tệp hỗ trợ: PNG, JPG, WEBP hoặc MP4.",
        "Video tối đa 15 giây; playlist tối đa 10 mục.",
        "Ảnh có thời lượng hiển thị 3 đến 15 giây tùy cấu hình từng mục.",
        "Có thể thêm tệp, sắp xếp playlist, điều chỉnh thời lượng ảnh và xóa mục.",
        "Bấm Lưu playlist để phát cấu hình mới sang màn hình khách.",
        "Cần cửa hàng đang chọn khớp cửa hàng đã gán thiết bị, quyền pos.advertising.read để xem và pos.advertising.manage để sửa.",
        "Nếu playlist trống, màn hình khách hiển thị hình Joy World mặc định.",
    ])
    add_steps(doc, [
        ("Bấm Thêm hình ảnh hoặc video.", "Chọn tệp đúng định dạng và giới hạn."),
        ("Kiểm tra mục vừa tải.", "Xem thumbnail, loại tệp và thời lượng."),
        ("Sắp xếp/thay đổi thời gian.", "Đặt thứ tự phát và thời gian hiển thị ảnh."),
        ("Bấm Lưu playlist.", "Chờ thông báo thành công rồi kiểm tra màn hình khách."),
    ])

    chapter_break(doc)
    add_heading(doc, "10. MÀN HÌNH KHÁCH HÀNG", 1)
    add_figure(doc, "15_man_hinh_khach_hang.png", "Màn hình khách ở trạng thái chờ - phát quảng cáo mặc định")
    add_heading(doc, "10.1. Các trạng thái hiển thị", 2)
    add_table(doc, ["Chế độ", "Khách hàng nhìn thấy"], [
        ["IDLE", "Quảng cáo/playlist; chưa có thông tin đơn."],
        ["CART", "Quảng cáo song song với danh sách hàng, số lượng, đơn giá, tổng tiền và thành viên (nếu có)."],
        ["TRANSFER", "Đơn chuyển khoản, QR, số tiền, nội dung, tài khoản và thời gian còn lại."],
        ["SUCCESS", "Xác nhận thanh toán thành công; QR cũ không còn hiển thị."],
        ["MEMBER_REVIEW", "Thông tin đăng ký/tra cứu thành viên để khách rà soát."],
        ["MEMBER_SUCCESS", "Xác nhận thành viên đã được tạo/cập nhật thành công."],
    ], [1800, 7560])
    add_heading(doc, "10.2. Cách kiểm tra", 2)
    add_steps(doc, [
        ("Ở trạng thái chờ.", "Kiểm tra quảng cáo chạy đúng, không màn hình đen."),
        ("Thêm một sản phẩm ở máy thu ngân.", "Màn hình khách phải chuyển sang CART và hiển thị đúng tổng tiền."),
        ("Tạo QR chuyển khoản.", "Màn hình khách phải hiển thị QR, số tiền và đếm ngược."),
        ("Hoàn tất giao dịch.", "Màn hình hiển thị thành công trong thời gian ngắn rồi trở về quảng cáo."),
        ("Thử chế độ thành viên.", "Thông tin khách cần kiểm tra phải khớp biểu mẫu/hồ sơ trên màn hình thu ngân."),
    ])
    add_note(doc, "Bảo vệ dữ liệu", "Hợp đồng dữ liệu màn hình khách không truyền UID nội bộ, token, bí mật, mã PayOS nội bộ, thông tin nhân viên hoặc trạng thái đồng bộ vận hành. Chỉ hiển thị dữ liệu cần thiết cho chính khách hàng.", "green")

    chapter_break(doc)
    add_heading(doc, "11. TRẠNG THÁI, CẢNH BÁO VÀ XỬ LÝ SỰ CỐ", 1)
    add_heading(doc, "11.1. Thiết bị và đăng nhập", 2)
    add_table(doc, ["Hiện tượng", "Cách xử lý"], [
        ["Mã kích hoạt không hợp lệ/hết hạn", "Yêu cầu mã 8 số mới trên JPULSE; nhập trong 10 phút và chỉ dùng một lần."],
        ["Thiết bị bị khóa/thu hồi", "Bấm Kiểm tra lại sau khi quản trị viên mở khóa; không xóa phiên chờ nếu chưa được hướng dẫn."],
        ["Quyền offline hết hạn", "Kết nối lại JPULSE để xác minh; sau 8 giờ JPOS không cho tiếp tục bằng cache cũ."],
        ["Đăng nhập sai", "Kiểm tra định danh/mật khẩu; dừng thử liên tục và liên hệ quản trị nếu tài khoản bị khóa."],
        ["Không có cửa hàng", "Quản trị viên cần gán cửa hàng loại STORE đang hoạt động cho tài khoản."],
    ], [3000, 6360])
    add_heading(doc, "11.2. Sản phẩm và giỏ hàng", 2)
    add_table(doc, ["Hiện tượng", "Cách xử lý"], [
        ["Không có sản phẩm", "Bấm Đồng bộ sản phẩm, kiểm tra mạng và catalog Firestore của cửa hàng."],
        ["Không tìm thấy mặt hàng", "Xóa từ khóa, chọn Tất cả/đúng loại, thử tên ngắn hơn hoặc mã vạch."],
        ["Giỏ hàng bị khóa", "Hoàn tất/hủy phiên QR PayOS đang hoạt động rồi mới sửa sản phẩm/thành viên."],
        ["Giá hoặc số vé bất thường", "Dừng bán mặt hàng, đồng bộ lại và báo quản trị catalog; không tự suy đoán giá."],
    ], [3000, 6360])
    add_heading(doc, "11.3. Thanh toán và đồng bộ", 2)
    add_table(doc, ["Hiện tượng", "Cách xử lý"], [
        ["QR đang tạo quá lâu", "Kiểm tra mạng, chờ thông báo lỗi; không bấm tạo nhiều phiên liên tiếp."],
        ["QR hết hạn", "Xác minh khách chưa chuyển tiền rồi tạo lại hoặc chọn phương thức khác."],
        ["QR cố định chưa xác nhận", "Kiểm tra giao dịch ngân hàng thực tế, sau đó xác nhận thủ công đúng một lần."],
        ["Đơn LOCAL_PAID/SYNCING", "Chờ tiến trình nền và làm mới lịch sử; không thu tiền lại."],
        ["Đơn SYNC_FAILED", "Mở chi tiết, đọc lỗi, đối chiếu thanh toán rồi bấm Thử lại."],
        ["Giao diện báo khôi phục", "Mở lịch sử theo mã local, xác minh trạng thái trước khi tiếp tục hoặc tạo thanh toán mới."],
    ], [3000, 6360])
    add_heading(doc, "11.4. In ấn", 2)
    add_bullets(doc, [
        "Không ra giấy: kiểm tra nguồn, cáp, giấy, hàng đợi in và máy in mặc định Windows.",
        "Bị cắt chữ: chọn đúng POS58/POS80/POS82 và đúng khổ trong driver.",
        "QR khó quét: tăng kích thước QR, giảm nhiễu logo/độ đậm và in thử.",
        "Vé không cắt riêng: bật cắt theo trang trong driver; mỗi vé là một trang.",
        "Thanh toán đã thành công nhưng in lỗi: in lại từ Chi tiết đơn; không thanh toán lại.",
    ])
    add_heading(doc, "11.5. Thành viên và quảng cáo", 2)
    add_bullets(doc, [
        "Đầu đọc thẻ lỗi: kiểm tra kết nối, hủy chờ, thử lại hoặc tra cứu bằng điện thoại.",
        "Không tải được hồ sơ/biến động/vé/thẻ: kiểm tra kích hoạt thiết bị, mạng và OpenAPI rồi bấm Làm mới/Thử lại.",
        "Không thấy Nạp bù: tài khoản chưa có quyền pos.members.compensate.",
        "Không sửa được quảng cáo: kiểm tra cửa hàng thiết bị, quyền read/manage và kết nối JPULSE.",
        "Tệp quảng cáo bị từ chối: dùng PNG/JPG/WEBP/MP4; video không quá 15 giây và playlist không quá 10 mục.",
    ])

    chapter_break(doc)
    add_heading(doc, "12. QUY TẮC VẬN HÀNH AN TOÀN VÀ PHỤ LỤC TRA CỨU NHANH", 1)
    add_heading(doc, "12.1. Quy tắc bắt buộc tại quầy", 2)
    add_bullets(doc, [
        "Luôn kiểm tra đúng cửa hàng, đúng người dùng và đúng ca trước khi bán.",
        "Đối chiếu sản phẩm, số lượng, thành viên, phương thức và tổng tiền trước khi xác nhận.",
        "Không thu lại tiền khi đơn chỉ đang chờ/lỗi đồng bộ hoặc khi máy in lỗi.",
        "Không xác nhận QR cố định nếu chưa nhìn thấy giao dịch thật trong tài khoản ngân hàng.",
        "Không tạo lại giao dịch khi chưa xác minh giao dịch cũ đã hết hạn/thất bại.",
        "Nạp bù/trừ điểm chỉ thực hiện khi có quyền và có lý do/chứng từ đã xác minh.",
        "Không chia sẻ mật khẩu, mã kích hoạt, token hay thông tin cấu hình nhạy cảm.",
        "Đăng xuất khi rời quầy hoặc bàn giao ca.",
    ])
    add_heading(doc, "12.2. Checklist mở ca", 2)
    add_bullets(doc, [
        "Thiết bị không bị khóa và quyền offline còn hiệu lực.",
        "Đăng nhập đúng tài khoản và cửa hàng.",
        "Đồng bộ sản phẩm; kiểm tra một số giá bán tiêu biểu.",
        "In thử biên lai và vé; QR trên bản in quét được.",
        "Đầu đọc thẻ hoạt động nếu quầy sử dụng.",
        "Màn hình khách hiển thị đúng quảng cáo và phản ánh giỏ hàng.",
        "Cấu hình PayOS/QR cố định đúng tài khoản của cửa hàng.",
    ])
    add_heading(doc, "12.3. Checklist đóng ca", 2)
    add_bullets(doc, [
        "Không còn khách đang chờ thanh toán hoặc QR đang hoạt động.",
        "Xử lý hoặc ghi nhận đầy đủ các đơn SYNC_FAILED/đang chờ.",
        "Tải báo cáo Kết ca đúng khoảng thời gian và phạm vi nhân viên.",
        "Đối chiếu tiền mặt, chuyển khoản và tổng doanh thu.",
        "In/lưu báo cáo theo quy trình nội bộ.",
        "Đăng xuất và bàn giao thiết bị, tiền mặt, biên lai/vé cho ca sau.",
    ])
    add_heading(doc, "12.4. Thông tin cần gửi khi yêu cầu hỗ trợ", 2)
    add_bullets(doc, [
        "Tên cửa hàng, tên quầy và thời điểm xảy ra lỗi.",
        "Tên màn hình/chức năng và các bước đã thực hiện.",
        "Mã local/mã HK của đơn (nếu liên quan); không gửi mật khẩu/token.",
        "Trạng thái đơn và thông báo lỗi nguyên văn.",
        "Phương thức thanh toán, số tiền và kết quả đối chiếu thực tế.",
        "Ảnh chụp màn hình lỗi và tình trạng máy in/đầu đọc/mạng.",
    ])
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
